from __future__ import annotations

import os
from pathlib import Path
import re
import shutil
import subprocess
import zipfile

import pytest
from docx import Document
from openpyxl import Workbook

from local_kb.extractors import base
from local_kb.extractors.base import (
    Extraction,
    Fragment,
    Registry,
    SnapshotCleanupError,
    registry,
)


def test_text_extractor_preserves_chinese_nonempty_line_locators(tmp_path: Path) -> None:
    path = tmp_path / "筆記.MD"
    path.write_text("第一行\n\n繁體中文內容\n", encoding="utf-8")

    result = registry.extract(path)

    assert result.status == "extracted"
    assert result.fragments == [
        Fragment("lines:1-1", "第一行"),
        Fragment("lines:3-3", "繁體中文內容"),
    ]


def test_unknown_media_is_pending_not_fabricated(tmp_path: Path) -> None:
    path = tmp_path / "recording.mp4"
    path.write_bytes(b"not-a-real-video")

    result = registry.extract(path)

    assert result.status == "pending_extractor"
    assert result.fragments == []
    assert result.warning == "no extractor for .mp4"


def test_unknown_format_validates_a_large_file_without_copying_a_snapshot(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = tmp_path / "large.mp4"
    path.write_bytes(b"x" * (2 * 1024 * 1024))
    copied = 0

    def copy_spy(fd: int, data: bytes) -> None:
        nonlocal copied
        copied += len(data)

    monkeypatch.setattr(base, "_write_all", copy_spy)

    result = registry.extract(path)

    assert result.status == "pending_extractor"
    assert result.fragments == []
    assert copied == 0


def test_suffix_matching_is_case_insensitive(tmp_path: Path) -> None:
    path = tmp_path / "UPPER.TXT"
    path.write_text("ok", encoding="utf-8")

    assert registry.extract(path).fragments == [Fragment("lines:1-1", "ok")]


def test_file_without_suffix_is_pending_with_clear_warning(tmp_path: Path) -> None:
    path = tmp_path / "README"
    path.write_text("untyped", encoding="utf-8")

    result = registry.extract(path)

    assert result.status == "pending_extractor"
    assert result.fragments == []
    assert result.warning == "no extractor for no suffix"


def test_html_strips_active_and_navigation_content_but_keeps_title_and_body(tmp_path: Path) -> None:
    path = tmp_path / "guide.html"
    path.write_text(
        "<html><head><title>指南</title><style>secret-style</style></head>"
        "<body><nav>secret-nav</nav><script>secret-script</script><p>正文</p></body></html>",
        encoding="utf-8",
    )

    result = registry.extract(path)

    assert result.status == "extracted"
    assert result.fragments == [Fragment("title:指南", "指南\n正文")]


def test_html_title_falls_back_to_filename_stem(tmp_path: Path) -> None:
    path = tmp_path / "untitled.htm"
    path.write_text("<p>body</p>", encoding="utf-8")

    assert registry.extract(path).fragments == [Fragment("title:untitled", "body")]


def test_pdf_with_text_creates_page_locators(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from local_kb.extractors import pdf

    path = tmp_path / "report.pdf"
    path.write_bytes(b"placeholder")

    class Page:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    class Reader:
        pages = [Page("首頁"), Page(""), Page("第二頁")]

    monkeypatch.setattr(pdf, "PdfReader", lambda _: Reader())

    assert registry.extract(path).fragments == [
        Fragment("page:1", "首頁"),
        Fragment("page:3", "第二頁"),
    ]


def test_blank_pdf_is_pending_for_ocr(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from local_kb.extractors import pdf

    path = tmp_path / "scan.pdf"
    path.write_bytes(b"placeholder")

    class Page:
        def extract_text(self) -> str:
            return "  "

    class Reader:
        pages = [Page()]

    monkeypatch.setattr(pdf, "PdfReader", lambda _: Reader())

    result = registry.extract(path)

    assert result.status == "pending_extractor"
    assert result.fragments == []
    assert result.warning == "PDF has no extractable text; OCR required"


def test_docx_uses_only_nonempty_paragraphs_with_paragraph_locators(tmp_path: Path) -> None:
    path = tmp_path / "memo.docx"
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("  ")
    document.add_paragraph("第三段")
    document.save(path)

    assert registry.extract(path).fragments == [
        Fragment("paragraph:1", "第一段"),
        Fragment("paragraph:3", "第三段"),
    ]


def test_docx_extracts_table_only_rows_and_deduplicates_merged_cells(tmp_path: Path) -> None:
    path = tmp_path / "table-only.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "名稱"
    table.cell(0, 1).text = "重要值"
    merged = table.cell(1, 0).merge(table.cell(1, 1))
    merged.text = "合併內容"
    document.save(path)

    result = registry.extract(path)

    assert result.fragments == [
        Fragment("table:1;row:1;cells:1-2", "名稱\t重要值"),
        Fragment("table:1;row:2;cells:1-1", "合併內容"),
    ]


def test_xlsx_extracts_each_nonempty_row_across_sheets_and_keeps_none_cells(tmp_path: Path) -> None:
    path = tmp_path / "table.xlsx"
    book = Workbook()
    first = book.active
    first.title = "摘要"
    first.append(["名稱", None, 3])
    first.append([None, None, None])
    second = book.create_sheet("資料")
    second.append([None, "值"])
    book.save(path)
    book.close()

    assert registry.extract(path).fragments == [
        Fragment("sheet:摘要;cells:A1-C1", "名稱\t\t3"),
        Fragment("sheet:資料;cells:A1-B1", "\t值"),
    ]


def test_xlsx_workbook_is_closed_after_extraction(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from local_kb.extractors import office

    path = tmp_path / "table.xlsx"
    real_book = Workbook()
    real_book.save(path)
    real_book.close()

    class Cell:
        def __init__(self, coordinate: str, value: object) -> None:
            self.coordinate = coordinate
            self.value = value

    class Sheet:
        title = "Data"
        max_row = 1
        max_column = 1

        def reset_dimensions(self) -> None:
            pass

        def iter_rows(self):
            yield (Cell("A1", "value"),)

    class Book:
        worksheets = [Sheet()]
        closed = False

        def close(self) -> None:
            self.closed = True

    book = Book()
    monkeypatch.setattr(office, "load_workbook", lambda *args, **kwargs: book)

    result = registry.extract(path)

    assert result.fragments == [Fragment("sheet:Data;cells:A1-A1", "value")]
    assert book.closed is True


@pytest.mark.parametrize("far_cell", ["A100001", "IW1", "IQ4000"])
def test_xlsx_rejects_sparse_dimensions_before_iterating_and_closes_book(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, far_cell: str
) -> None:
    from local_kb.extractors import office

    path = tmp_path / "sparse.xlsx"
    source_book = Workbook()
    source_book.active[far_cell] = "far away"
    source_book.save(path)
    source_book.close()
    real_load_workbook = office.load_workbook
    loaded: list[object] = []

    class SheetProxy:
        def __init__(self, sheet) -> None:
            self.title = sheet.title
            self.max_row = sheet.max_row
            self.max_column = sheet.max_column

        def reset_dimensions(self) -> None:
            raise AssertionError("declared dimension precheck must run before reset")

        def iter_rows(self):
            raise AssertionError("dimension budget must be checked before iter_rows")

    class BookProxy:
        def __init__(self, book) -> None:
            self.book = book
            self.worksheets = [SheetProxy(sheet) for sheet in book.worksheets]
            self.closed = False

        def close(self) -> None:
            self.closed = True
            self.book.close()

    def tracked_load(*args, **kwargs):
        proxy = BookProxy(real_load_workbook(*args, **kwargs))
        loaded.append(proxy)
        return proxy

    monkeypatch.setattr(office, "load_workbook", tracked_load)

    with pytest.raises(base.ExtractionError, match="worksheet dimensions exceed budget"):
        registry.extract(path)

    assert len(loaded) == 1
    assert loaded[0].closed is True


def test_xlsx_reset_dimensions_recovers_a_real_cell_hidden_by_low_dimension(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.xlsx"
    path = tmp_path / "low-dimension.xlsx"
    book = Workbook()
    book.active["B2"] = "hidden-by-dimension"
    book.save(source)
    book.close()

    with zipfile.ZipFile(source) as original, zipfile.ZipFile(
        path, "w", compression=zipfile.ZIP_DEFLATED
    ) as rewritten:
        for member in original.infolist():
            data = original.read(member)
            if member.filename == "xl/worksheets/sheet1.xml":
                data = re.sub(
                    br'<dimension ref="[^"]+"',
                    b'<dimension ref="A1:A1"',
                    data,
                    count=1,
                )
            rewritten.writestr(member, data)

    result = registry.extract(path)

    assert result.fragments == [
        Fragment("sheet:Sheet;cells:A2-B2", "\thidden-by-dimension")
    ]


@pytest.mark.parametrize("actual_shape", ["far_row", "far_column", "cell_product"])
def test_xlsx_low_dimension_rejects_actual_bounds_before_cell_stringification(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, actual_shape: str
) -> None:
    from local_kb.extractors import office

    path = tmp_path / "low-dimension.xlsx"
    source_book = Workbook()
    source_book.save(path)
    source_book.close()
    value_reads = 0
    rows_yielded = 0

    class ForbiddenCell:
        @property
        def value(self):
            nonlocal value_reads
            value_reads += 1
            raise AssertionError("budget must be checked before reading cell values")

    class Sheet:
        title = "Low"
        max_row = 1
        max_column = 1
        reset_called = False

        def reset_dimensions(self) -> None:
            self.reset_called = True

        def iter_rows(self):
            nonlocal rows_yielded
            if actual_shape == "far_column":
                rows_yielded += 1
                yield tuple(ForbiddenCell() for _ in range(257))
                return
            if actual_shape == "cell_product":
                for _ in range(3_999):
                    rows_yielded += 1
                    yield ()
                rows_yielded += 1
                yield tuple(ForbiddenCell() for _ in range(251))
                return
            for _ in range(100_000):
                rows_yielded += 1
                yield ()
            rows_yielded += 1
            yield (ForbiddenCell(),)

    class Book:
        def __init__(self) -> None:
            self.worksheets = [Sheet()]
            self.closed = False

        def close(self) -> None:
            self.closed = True

    loaded = Book()
    monkeypatch.setattr(office, "load_workbook", lambda *args, **kwargs: loaded)

    with pytest.raises(base.ExtractionError, match="worksheet dimensions exceed budget"):
        registry.extract(path)

    assert loaded.worksheets[0].reset_called is True
    assert value_reads == 0
    expected_rows = {"far_column": 1, "cell_product": 4_000, "far_row": 100_001}
    assert rows_yielded == expected_rows[actual_shape]
    assert loaded.closed is True


@pytest.mark.parametrize("suffix", [".pdf", ".docx", ".xlsx"])
def test_corrupt_supported_documents_never_claim_extracted(tmp_path: Path, suffix: str) -> None:
    path = tmp_path / f"corrupt{suffix}"
    path.write_bytes(b"not a document")

    with pytest.raises(Exception):
        registry.extract(path)


def test_registry_rejects_duplicate_suffixes() -> None:
    class One:
        suffixes = {".one"}

        def extract(self, path: Path) -> Extraction:
            return Extraction("extracted", [])

    class Two:
        suffixes = {".ONE"}

        def extract(self, path: Path) -> Extraction:
            return Extraction("extracted", [])

    isolated = Registry()
    isolated.register(One())

    with pytest.raises(ValueError, match="duplicate extractor suffix: .one"):
        isolated.register(Two())


def test_extractions_do_not_share_mutable_fragment_lists() -> None:
    first = Extraction("extracted", [])
    second = Extraction("extracted", [])

    first.fragments.append(Fragment("lines:1-1", "only first"))

    assert second.fragments == []


def test_registry_rejects_directories_and_symlinks(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="regular file"):
        registry.extract(tmp_path)

    target = tmp_path / "target.txt"
    target.write_text("not read by link", encoding="utf-8")
    link = tmp_path / "link.txt"
    try:
        link.symlink_to(target)
    except OSError:
        pytest.skip("symlink creation is unavailable in this environment")

    with pytest.raises(ValueError, match="symbolic link"):
        registry.extract(link)


def test_registry_rejects_a_file_below_a_parent_junction(tmp_path: Path) -> None:
    if os.name != "nt":
        pytest.skip("junctions are a Windows-only filesystem feature")
    outside = tmp_path / "outside"
    outside.mkdir()
    (outside / "secret.mp4").write_text("must not be read", encoding="utf-8")
    junction = tmp_path / "linked-parent"
    created = subprocess.run(
        ["cmd.exe", "/d", "/c", "mklink /J linked-parent outside"],
        cwd=tmp_path,
        check=False,
        capture_output=True,
        text=True,
    )
    assert created.returncode == 0, created.stderr

    with pytest.raises(ValueError, match="reparse"):
        registry.extract(junction / "secret.mp4")


def test_registry_extracts_a_stable_snapshot_not_a_path_reopened_after_dispatch(
    tmp_path: Path,
) -> None:
    source = tmp_path / "mutable.snap"
    source.write_text("before", encoding="utf-8")

    class ReplacingExtractor:
        suffixes = {".snap"}

        def extract(self, path: Path) -> Extraction:
            source.write_text("after", encoding="utf-8")
            return Extraction("extracted", [Fragment("snapshot", path.read_text(encoding="utf-8"))])

    isolated = Registry()
    isolated.register(ReplacingExtractor())

    result = isolated.extract(source)

    assert result.fragments == [Fragment("snapshot", "before")]
    assert source.read_text(encoding="utf-8") == "after"


def test_parent_replacement_after_snapshot_cannot_change_parser_input(tmp_path: Path) -> None:
    parent = tmp_path / "source-parent"
    parent.mkdir()
    source = parent / "race.snap"
    source.write_text("before", encoding="utf-8")
    replacement = tmp_path / "replaced-parent"
    rename_failed = False

    class ParentReplacingExtractor:
        suffixes = {".snap"}

        def extract(self, path: Path) -> Extraction:
            nonlocal rename_failed
            try:
                parent.rename(replacement)
                parent.mkdir()
                (parent / "race.snap").write_text("outside", encoding="utf-8")
            except OSError:
                rename_failed = True
            return Extraction("extracted", [Fragment("snapshot", path.read_text(encoding="utf-8"))])

    isolated = Registry()
    isolated.register(ParentReplacingExtractor())

    result = isolated.extract(source)

    assert result.fragments == [Fragment("snapshot", "before")]
    if os.name == "nt":
        assert rename_failed is True


def test_registry_removes_snapshot_after_success_and_parser_failure(tmp_path: Path) -> None:
    source = tmp_path / "temporary.snap"
    source.write_text("content", encoding="utf-8")
    seen_paths: list[Path] = []

    class CapturingExtractor:
        suffixes = {".snap"}

        def extract(self, path: Path) -> Extraction:
            seen_paths.append(path)
            assert path.exists()
            return Extraction("extracted", [])

    isolated = Registry()
    isolated.register(CapturingExtractor())
    isolated.extract(source)
    assert not seen_paths[0].exists()
    assert not seen_paths[0].parent.exists()

    class FailingExtractor:
        suffixes = {".fail"}

        def extract(self, path: Path) -> Extraction:
            seen_paths.append(path)
            raise RuntimeError("parser failed")

    failing_source = tmp_path / "temporary.fail"
    failing_source.write_text("content", encoding="utf-8")
    failing = Registry()
    failing.register(FailingExtractor())
    with pytest.raises(RuntimeError, match="parser failed"):
        failing.extract(failing_source)
    assert not seen_paths[1].exists()
    assert not seen_paths[1].parent.exists()


def test_registry_raises_if_a_parser_leaves_its_snapshot_open(tmp_path: Path) -> None:
    source = tmp_path / "held.open"
    source.write_text("content", encoding="utf-8")
    held = None

    class HandleHoldingExtractor:
        suffixes = {".open"}

        def extract(self, path: Path) -> Extraction:
            nonlocal held
            held = path.open("rb")
            return Extraction("extracted", [Fragment("snapshot", "content")])

    isolated = Registry()
    isolated.register(HandleHoldingExtractor())

    with pytest.raises(SnapshotCleanupError) as error:
        isolated.extract(source)

    assert error.value.snapshot_directory.exists()
    assert held is not None
    held.close()
    shutil.rmtree(error.value.snapshot_directory)
    assert not error.value.snapshot_directory.exists()


def test_supported_raw_source_budget_is_exact_and_unknown_files_are_not_limited(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    assert base.MAX_SOURCE_BYTES == 100 * 1024 * 1024
    monkeypatch.setattr(base, "MAX_SOURCE_BYTES", 16)
    exact = tmp_path / "exact.txt"
    exact.write_bytes(b"x" * 16)
    oversized = tmp_path / "oversized.txt"
    oversized.write_bytes(b"x" * 17)
    unknown = tmp_path / "oversized.mp4"
    unknown.write_bytes(b"x" * 17)

    assert registry.extract(exact).status == "extracted"
    with pytest.raises(base.ExtractionError, match="source exceeds 100 MiB budget"):
        registry.extract(oversized)
    renamed = tmp_path / "released.txt"
    oversized.rename(renamed)
    assert renamed.exists()
    assert registry.extract(unknown).status == "pending_extractor"


def test_high_compression_docx_is_rejected_from_zip_metadata(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import office

    path = tmp_path / "bomb.docx"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"x" * (11 * 1024 * 1024))

    snapshots: list[Path] = []
    validate_zip = office._validate_office_zip

    def tracked_validation(snapshot: Path) -> None:
        snapshots.append(snapshot)
        validate_zip(snapshot)

    monkeypatch.setattr(office, "_validate_office_zip", tracked_validation)

    assert path.stat().st_size < 100_000
    with pytest.raises(base.ExtractionError, match="compression ratio exceeds budget"):
        registry.extract(path)
    assert not snapshots[0].exists()
    assert not snapshots[0].parent.exists()


def test_single_large_office_xml_is_rejected_before_parser(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import office

    monkeypatch.setattr(base, "MAX_ZIP_SINGLE_XML_BYTES", 64)
    path = tmp_path / "expanded.docx"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
        archive.writestr("word/document.xml", b"x" * 65)
    parser_called = False

    def parser_spy(path):
        nonlocal parser_called
        parser_called = True
        raise AssertionError("Document parser must not run")

    monkeypatch.setattr(office, "Document", parser_spy)

    with pytest.raises(base.ExtractionError, match="single XML member exceeds budget"):
        registry.extract(path)

    assert parser_called is False


@pytest.mark.parametrize(
    ("limit_name", "fragments", "message"),
    [
        ("MAX_FRAGMENT_COUNT", [Fragment("1", "a"), Fragment("2", "b")], "fragment count"),
        ("MAX_EXTRACTION_CHARACTERS", [Fragment("1", "abc")], "character count"),
    ],
)
def test_registry_rejects_over_budget_extractor_output_and_cleans_snapshot(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    limit_name: str,
    fragments: list[Fragment],
    message: str,
) -> None:
    source = tmp_path / "output.budget"
    source.write_text("source", encoding="utf-8")
    monkeypatch.setattr(base, limit_name, 1 if limit_name == "MAX_FRAGMENT_COUNT" else 2)
    snapshots: list[Path] = []

    class OversizedOutputExtractor:
        suffixes = {".budget"}

        def extract(self, path: Path) -> Extraction:
            snapshots.append(path)
            return Extraction("extracted", fragments)

    isolated = Registry()
    isolated.register(OversizedOutputExtractor())

    with pytest.raises(base.ExtractionError, match=message):
        isolated.extract(source)

    assert not snapshots[0].exists()
    assert not snapshots[0].parent.exists()


def test_direct_text_extractor_enforces_output_budget_and_cleans_snapshot(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import text

    source = tmp_path / "direct.txt"
    source.write_text("source", encoding="utf-8")
    monkeypatch.setattr(base, "MAX_FRAGMENT_COUNT", 1)
    snapshots: list[Path] = []

    def oversized_output(path: Path) -> Extraction:
        snapshots.append(path)
        return Extraction("extracted", [Fragment("1", "a"), Fragment("2", "b")])

    monkeypatch.setattr(text, "_extract_text_snapshot", oversized_output)

    with pytest.raises(base.ExtractionError, match="fragment count"):
        text.TextExtractor().extract(source)

    assert not snapshots[0].exists()
    assert not snapshots[0].parent.exists()


def test_snapshot_copy_stops_when_a_small_source_grows_over_budget(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "growing.txt"
    source.write_bytes(b"x")
    monkeypatch.setattr(base, "MAX_SOURCE_BYTES", 1)
    chunks = iter([b"x", b"y"])
    monkeypatch.setattr(base.os, "read", lambda fd, size: next(chunks, b""))
    created: list[Path] = []
    real_mkdtemp = base.tempfile.mkdtemp

    def tracked_mkdtemp(*args, **kwargs):
        directory = Path(real_mkdtemp(*args, **kwargs))
        created.append(directory)
        return str(directory)

    monkeypatch.setattr(base.tempfile, "mkdtemp", tracked_mkdtemp)

    with pytest.raises(base.ExtractionError, match="grew beyond source budget"):
        registry.extract(source)

    assert len(created) == 1
    assert not created[0].exists()


def test_snapshot_copy_rejects_source_metadata_change_and_cleans_temp(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "changing.txt"
    source.write_bytes(b"abcdef")
    metadata = iter([(6, 100), (5, 101)])
    monkeypatch.setattr(base, "_source_metadata", lambda fd: next(metadata))
    created: list[Path] = []
    real_mkdtemp = base.tempfile.mkdtemp

    def tracked_mkdtemp(*args, **kwargs):
        directory = Path(real_mkdtemp(*args, **kwargs))
        created.append(directory)
        return str(directory)

    monkeypatch.setattr(base.tempfile, "mkdtemp", tracked_mkdtemp)

    with pytest.raises(base.ExtractionError, match="source changed during snapshot copy"):
        registry.extract(source)

    assert len(created) == 1
    assert not created[0].exists()


def test_pdf_fragment_budget_stops_page_generator_at_second_fragment(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import pdf

    source = tmp_path / "many.pdf"
    source.write_bytes(b"placeholder")
    monkeypatch.setattr(base, "MAX_FRAGMENT_COUNT", 1)
    pages_visited = 0
    snapshots: list[Path] = []

    class Page:
        def extract_text(self) -> str:
            return "text"

    class Pages:
        def __iter__(self):
            nonlocal pages_visited
            for _ in range(10_000):
                pages_visited += 1
                yield Page()

    class Reader:
        pages = Pages()

    monkeypatch.setattr(pdf, "PdfReader", lambda stream: Reader())
    validate_pdf = pdf._extract_pdf_snapshot

    def tracked_extract(path: Path) -> Extraction:
        snapshots.append(path)
        return validate_pdf(path)

    monkeypatch.setattr(pdf.PdfExtractor, "extract_snapshot", staticmethod(tracked_extract))

    with pytest.raises(base.ExtractionError, match="fragment count"):
        registry.extract(source)

    assert pages_visited == 2
    assert not snapshots[0].exists()
    assert not snapshots[0].parent.exists()


def test_html_rejects_single_oversized_body_before_fragment_construction(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import html

    source = tmp_path / "large.html"
    source.write_text("<p>ab</p>", encoding="utf-8")
    monkeypatch.setattr(base, "MAX_EXTRACTION_CHARACTERS", 1)
    constructed = 0
    real_fragment = base.Fragment

    def fragment_spy(locator: str, text: str) -> Fragment:
        nonlocal constructed
        constructed += 1
        return real_fragment(locator, text)

    monkeypatch.setattr(base, "Fragment", fragment_spy)

    with pytest.raises(base.ExtractionError, match="character count"):
        html.HtmlExtractor().extract(source)

    assert constructed == 0


def test_text_extractor_streams_lines_without_read_text(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import text

    source = tmp_path / "stream.txt"
    source.write_text("first\nsecond\n", encoding="utf-8")

    def forbidden_read_text(*args, **kwargs):
        raise AssertionError("text extractor must stream lines")

    monkeypatch.setattr(Path, "read_text", forbidden_read_text)

    result = text.TextExtractor().extract(source)

    assert result.fragments == [
        Fragment("lines:1-1", "first"),
        Fragment("lines:2-2", "second"),
    ]


@pytest.mark.parametrize(
    "separator",
    ["\n", "\r", "\r\n", "\v", "\f", "\x1c", "\x1d", "\x1e", "\x85", "\u2028", "\u2029"],
)
def test_text_streaming_supports_every_splitlines_separator(
    tmp_path: Path, separator: str
) -> None:
    path = tmp_path / "separator.txt"
    path.write_text(f"first{separator}second", encoding="utf-8", newline="")

    result = registry.extract(path)

    assert result.fragments == [
        Fragment("lines:1-1", "first"),
        Fragment("lines:2-2", "second"),
    ]


def test_text_streaming_handles_mixed_unicode_boundaries(tmp_path: Path) -> None:
    path = tmp_path / "mixed.txt"
    path.write_text("first\u2028second\x85third\vfourth\fform", encoding="utf-8")

    assert registry.extract(path).fragments == [
        Fragment("lines:1-1", "first"),
        Fragment("lines:2-2", "second"),
        Fragment("lines:3-3", "third"),
        Fragment("lines:4-4", "fourth"),
        Fragment("lines:5-5", "form"),
    ]


def test_text_streaming_treats_cross_chunk_crlf_as_one_boundary(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import text

    path = tmp_path / "crlf.txt"
    path.write_bytes(b"first\r\nsecond")
    monkeypatch.setattr(text, "TEXT_CHUNK_CHARACTERS", 6)

    assert registry.extract(path).fragments == [
        Fragment("lines:1-1", "first"),
        Fragment("lines:2-2", "second"),
    ]


def test_text_streaming_counts_ignored_blank_logical_lines(tmp_path: Path) -> None:
    path = tmp_path / "blank-lines.txt"
    path.write_text("first\u2028 \x85third", encoding="utf-8")

    assert registry.extract(path).fragments == [
        Fragment("lines:1-1", "first"),
        Fragment("lines:3-3", "third"),
    ]


def test_text_streaming_rejects_overlong_unbroken_line_before_reading_more(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from local_kb.extractors import text

    source = tmp_path / "overlong.txt"
    source.write_text("placeholder", encoding="utf-8")
    monkeypatch.setattr(base, "MAX_EXTRACTION_CHARACTERS", 1)
    monkeypatch.setattr(text, "TEXT_CHUNK_CHARACTERS", 2)
    reads = 0

    class FakeStream:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self, size: int) -> str:
            nonlocal reads
            reads += 1
            return "ab" if reads == 1 else "should-not-be-read"

    monkeypatch.setattr(Path, "open", lambda *args, **kwargs: FakeStream())

    with pytest.raises(base.ExtractionError, match="unbroken line exceeds"):
        text.TextExtractor().extract(source)

    assert reads == 1
